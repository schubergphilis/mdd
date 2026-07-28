"""docx.py — DocxConverter: convert .docx files to Markdown.

Conversion logic moved from mdd.commands.convert; behaviour is unchanged.
"""

import html
import logging
import re
from collections import Counter
from typing import TYPE_CHECKING, Any

from mdd.convert import CorruptSourceError
from mdd.converters.protocol import ConvertResult
from mdd.utils.logging import get_logger

log = get_logger(__name__)

if TYPE_CHECKING:
    from collections.abc import Callable, Iterator
    from pathlib import Path
    from types import TracebackType

    from docling.document_converter import (
        DocumentConverter,  # pyright: ignore[reportMissingModuleSource]
    )

# ---------------------------------------------------------------------------
# Constants shared with the old commands/convert module (kept for parity)
# ---------------------------------------------------------------------------

TOC_HEADING_RE = re.compile(
    r"^(?:#{7,}\s.*|#{1,6}\s+(?:table\s+of\s+contents|contents))\s*$", re.IGNORECASE
)
TOC_LINE_RE = re.compile(r"^\s*(\d+\.|\*|-)\s+.+\d+\s*$")
TOC_LINE_TAB_RE = re.compile(r"^\d+(\.\d+)*\t.+\t\d+\s*$")

_KNOWN_META_KEYS = {
    "date",
    "location",
    "participants",
    "invitees",
    "invited",
    "invites",
    "absent",
    "apologies",
    "for",
    "author",
    "status",
}
_KEY_NORMALISE = {"invited": "Invitees", "apologies": "Absent"}
_MULTI_VALUE_KEYS = {"Participants", "Invitees", "Invites", "Absent"}
_DOC_META_KEYS = {"For", "Date", "Author", "Status"}

_BOILERPLATE_SECTION_RE = re.compile(
    r"^(#{1,6})\s+(?:\d+\.?\s+)?(meeting\s+logistics)\s*$", re.IGNORECASE
)
_BOILERPLATE_LINE_RES = [
    re.compile(r"^\s*[-*]\s+Actions are on the \[", re.IGNORECASE),
]

# site-specific DOCX heuristics (spec S44 / plan P03 MR A5): the two-column
# metadata-table extraction (meeting-minutes "For / Date / Author / …" tables)
# and the "Meeting Logistics" boilerplate stripper. On by default so the
# original behaviour is unchanged; a deployment without those templates
# turns it off via
# :func:`set_sbp_docx_heuristics` (the open-source core defaults it off in the
# reference wrapper — plan P03 phases C–E).
_sbp_docx_heuristics = True


def set_sbp_docx_heuristics(*, enabled: bool) -> None:
    """Enable or disable the site-specific DOCX metadata + boilerplate heuristics."""
    global _sbp_docx_heuristics
    _sbp_docx_heuristics = enabled


_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

_converter: DocumentConverter | None = None


def _get_converter() -> DocumentConverter:
    global _converter
    if _converter is None:
        # lazy: docling pulls torch/transformers; loaded only on docx conversion
        from docling.document_converter import (  # noqa: PLC0415
            DocumentConverter,  # pyright: ignore[reportMissingModuleSource]
        )

        _converter = DocumentConverter()
    return _converter  # pyright: ignore[reportReturnType]


def _all_tables(doc: Any) -> Iterator[Any]:  # pyright: ignore[reportAny]
    """Yield all tables depth-first, innermost first."""

    def _recurse(table: Any) -> Iterator[Any]:  # pyright: ignore[reportAny]
        for row in table.rows:
            for cell in row.cells:
                for nested in cell.tables:
                    yield from _recurse(nested)
        yield table

    for top in doc.tables:
        yield from _recurse(top)


def _cell_full_text(cell: Any) -> str:  # pyright: ignore[reportAny]
    """Get cell text including content inside SDT elements."""
    return "".join(t.text for t in cell._element.findall(f".//{{{_W}}}t") if t.text).strip()


def _cell_para_texts(cell: Any) -> list[str]:  # pyright: ignore[reportAny]
    """Per-paragraph texts from a cell, skipping empty paragraphs."""
    result: list[str] = []
    for para in cell.paragraphs:
        text = "".join(t.text for t in para._element.findall(f".//{{{_W}}}t") if t.text).strip()
        if text:
            result.append(text)
    return result


def _is_meta_table(table: Any) -> bool:  # pyright: ignore[reportAny]
    """A metadata table is 2-column, ≤ 20 rows, with at least one known key."""
    if len(table.columns) != 2 or len(table.rows) > 20:  # pyright: ignore[reportAny]
        return False
    first_col = (_cell_full_text(r.cells[0]).lower() for r in table.rows if r.cells)  # pyright: ignore[reportAny]
    return any(v in _KNOWN_META_KEYS for v in first_col)


def _row_meta_entry(row: Any) -> tuple[str, str] | None:  # pyright: ignore[reportAny]
    """Return ``(normalised_key, value)`` for a metadata row, or None to skip."""
    raw_key = _cell_full_text(row.cells[0]).lower()  # pyright: ignore[reportAny]
    if raw_key not in _KNOWN_META_KEYS:
        return None
    normalised: str = _KEY_NORMALISE.get(raw_key) or raw_key.capitalize()
    paras = _cell_para_texts(row.cells[1])  # pyright: ignore[reportAny]
    value = ", ".join(paras) if paras else _cell_full_text(row.cells[1])  # pyright: ignore[reportAny]
    if not value:
        return None
    return normalised, value


def extract_metadata(path: Path) -> dict[str, str]:
    """Extract house-template metadata from all matching two-column tables.

    Returns ``{}`` when the house-template DOCX heuristics are disabled.
    """
    if not _sbp_docx_heuristics:
        return {}
    # lazy: python-docx pulls lxml; loaded only when reading docx
    from docx import Document  # pyright: ignore[reportMissingModuleSource]  # noqa: PLC0415

    doc: Any = Document(str(path))  # pyright: ignore[reportAny]
    result: dict[str, str] = {}
    for table in _all_tables(doc):
        if not _is_meta_table(table):
            continue
        for row in table.rows:  # pyright: ignore[reportAny]
            entry = _row_meta_entry(row)
            if entry is None or entry[0] in result:
                continue
            result[entry[0]] = entry[1]
    return result


def _para_run_texts(element: Any) -> list[str]:  # pyright: ignore[reportAny]
    """Return one string per <w:p> paragraph, joining all <w:t> runs."""
    result: list[str] = []
    for para in element.findall(f".//{{{_W}}}p"):
        text = "".join(t.text or "" for t in para.findall(f".//{{{_W}}}t")).strip()
        if text:
            result.append(text)
    return result


def _title_from_core_properties(doc: Any) -> str:  # pyright: ignore[reportAny]
    return (doc.core_properties.title or "").strip()  # pyright: ignore[reportAny]


def _title_from_document_title_sdt(doc: Any) -> str:  # pyright: ignore[reportAny]
    """Return text of the SDT whose alias is ``DocumentTitle`` (first match)."""
    for sdt in doc.element.findall(f".//{{{_W}}}sdt"):  # pyright: ignore[reportAny]
        alias = sdt.find(f"{{{_W}}}sdtPr/{{{_W}}}alias")
        if alias is not None and alias.get(f"{{{_W}}}val") == "DocumentTitle":
            texts = _para_run_texts(sdt)
            if texts:
                return " ".join(texts)
    return ""


def _title_from_cover_title_paragraphs(doc: Any) -> str:  # pyright: ignore[reportAny]
    """Concatenate text of all paragraphs styled ``Cover-Title``."""
    titles: list[str] = []
    for para in doc.element.findall(f".//{{{_W}}}p"):  # pyright: ignore[reportAny]
        pstyle = para.find(f"{{{_W}}}pPr/{{{_W}}}pStyle")
        if pstyle is None or pstyle.get(f"{{{_W}}}val") != "Cover-Title":
            continue
        text = "".join(t.text or "" for t in para.findall(f".//{{{_W}}}t")).strip()
        if text:
            titles.append(text)
    return " ".join(titles)


# Title sources are tried in priority order; first non-empty wins. Each source
# returns "" when it has no opinion, so the dispatch loop stays a simple chain
# rather than three nested branches in extract_title.
_TITLE_SOURCES: tuple[Callable[[Any], str], ...] = (
    _title_from_core_properties,
    _title_from_document_title_sdt,
    _title_from_cover_title_paragraphs,
)


def extract_title(path: Path) -> str:
    """Extract document title.

    Checks: core_properties.title, DocumentTitle SDT alias, Cover-Title paragraph style.
    """
    # lazy: python-docx pulls lxml; loaded only when reading docx
    from docx import Document  # pyright: ignore[reportMissingModuleSource]  # noqa: PLC0415

    doc: Any = Document(str(path))  # pyright: ignore[reportAny]
    for source in _TITLE_SOURCES:
        if title := source(doc):
            return title
    return ""


def strip_empty_headings(lines: list[str]) -> list[str]:
    """Remove empty valid headings (##) and invalid artifact headings (7+ #)."""
    return [
        line
        for line in lines
        if not re.match(r"^#{1,6}\s*$", line) and not re.match(r"^#{7,}", line)
    ]


def _is_toc_line(line: str) -> bool:
    return bool(TOC_LINE_RE.match(line) or TOC_LINE_TAB_RE.match(line))


def strip_toc(lines: list[str]) -> list[str]:
    """Strip Word-generated table-of-contents blocks."""
    result: list[str] = []
    i = 0
    while i < len(lines):
        if TOC_HEADING_RE.match(lines[i]):
            lookahead = [ln for ln in lines[i + 1 : i + 10] if ln.strip()][:3]
            if any(_is_toc_line(ln) for ln in lookahead):
                i += 1
                while i < len(lines) and (not lines[i].strip() or _is_toc_line(lines[i])):
                    i += 1
                continue
        result.append(lines[i])
        i += 1
    return result


def strip_leading_table(md: str) -> str:
    """Remove the first markdown table block and its trailing blank lines."""
    lines = md.splitlines()
    i = 0
    while i < len(lines) and lines[i].startswith("|"):
        i += 1
    if i > 0:
        while i < len(lines) and not lines[i].strip():
            i += 1
    return "\n".join(lines[i:])


def strip_boilerplate_sections(lines: list[str]) -> list[str]:
    """Strip entire sections whose heading matches Meeting Logistics pattern."""
    result: list[str] = []
    i = 0
    while i < len(lines):
        m = _BOILERPLATE_SECTION_RE.match(lines[i])
        if m:
            level = len(m.group(1))
            i += 1
            while i < len(lines) and not re.match(rf"^#{{1,{level}}}\s", lines[i]):
                i += 1
            continue
        result.append(lines[i])
        i += 1
    return result


def strip_boilerplate_lines(lines: list[str]) -> list[str]:
    """Strip individual boilerplate lines."""
    return [line for line in lines if not any(p.match(line) for p in _BOILERPLATE_LINE_RES)]


def _postprocess(md: str) -> str:
    md = html.unescape(md)
    lines = md.splitlines()
    lines = strip_toc(lines)
    lines = strip_empty_headings(lines)
    if _sbp_docx_heuristics:
        lines = strip_boilerplate_sections(lines)
        lines = strip_boilerplate_lines(lines)
    return "\n".join(lines)


_DOCLING_MSWORD_LOGGER = "docling.backend.msword_backend"
_DROPPED_IMAGE_PATTERNS = (
    "image cannot be loaded by Pillow",
    "DrawingML image cannot be loaded by Pillow",
    "VML image cannot be loaded",
)


class _DroppedImageCounter(logging.Filter):
    """Filter that swallows Docling's per-image Pillow warnings and counts them.

    Docling logs one warning per embedded image it cannot rasterize (typically
    WMF/EMF metafiles in Word templates), and the default message includes
    only ``<_io.BytesIO object at 0x...>`` which is unactionable. We suppress
    those records and emit a single end-of-conversion summary instead.
    """

    def __init__(self) -> None:
        super().__init__()
        self.dropped = 0

    def filter(self, record: logging.LogRecord) -> bool:
        msg = record.getMessage()
        if any(p in msg for p in _DROPPED_IMAGE_PATTERNS):
            self.dropped += 1
            return False
        return True


class _SuppressDroppedImageWarnings:
    """Context manager that installs :class:`_DroppedImageCounter` on docling.

    On exit, emits a single summary line to stderr if any image warnings
    were swallowed — replacing N noisy per-image log lines with one useful
    message naming the source file.
    """

    def __init__(self, source: Path | None = None) -> None:
        self.counter = _DroppedImageCounter()
        self._source = source

    def __enter__(self) -> _DroppedImageCounter:
        logging.getLogger(_DOCLING_MSWORD_LOGGER).addFilter(self.counter)
        return self.counter

    def __exit__(
        self,
        exc_type: type[BaseException] | None,
        exc: BaseException | None,
        tb: TracebackType | None,
    ) -> None:
        logging.getLogger(_DOCLING_MSWORD_LOGGER).removeFilter(self.counter)
        if self.counter.dropped and exc_type is None:
            name = self._source.name if self._source is not None else "docx"
            log.warning(
                "%s: dropped %d embedded image(s) that could not be loaded "
                "(typically WMF/EMF metafiles). "
                "Install LibreOffice for better WMF/EMF support.",
                name,
                self.counter.dropped,
            )


def convert_body(path: Path) -> str:
    """Convert .docx body to Markdown using Docling."""
    with _SuppressDroppedImageWarnings(path):
        result = _get_converter().convert(str(path))
    md: str = result.document.export_to_markdown()
    return _postprocess(md)


# DrawingML embed reference: <a:blip r:embed="rIdN"> identifies an image
# part via the document-part relationship table. We collect these in
# document order so the extracted images line up 1:1 with the
# <!-- image --> placeholders Docling emits in the markdown.
_A = "http://schemas.openxmlformats.org/drawingml/2006/main"
_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_IMAGE_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"

_CONTENT_TYPE_TO_FORMAT: dict[str, str] = {
    "image/png": "png",
    "image/jpeg": "jpeg",
    "image/jpg": "jpeg",
    "image/gif": "gif",
    "image/tiff": "tiff",
    "image/x-tiff": "tiff",
    "image/x-wmf": "wmf",
    "image/x-emf": "emf",
    "image/bmp": "bmp",
}


def _ordered_image_rids(doc: Any) -> list[str]:  # pyright: ignore[reportAny]
    """Return the relationship IDs of embedded images in document order."""
    body: Any = doc.element.body  # pyright: ignore[reportAny]
    rids: list[str] = []
    for blip in body.iter(f"{{{_A}}}blip"):  # pyright: ignore[reportAny]
        rid: str | None = blip.get(f"{{{_R}}}embed")  # pyright: ignore[reportAny]
        if rid:
            rids.append(rid)
    return rids


def _extract_docx_images(
    src: Path,
    attachments_dir: Path,
    dropped: Counter[str],
) -> list[str]:
    """Extract embedded images from *src* into *attachments_dir*.

    Returns one markdown link per ``<a:blip>`` reference in document order.
    Unrecognised content types are dropped via *dropped* and yield an empty
    placeholder marker — the caller injects ``"<!-- image dropped: <fmt> -->"``
    in that position so the slot in the markdown is preserved.
    """
    # lazy: python-docx pulls lxml; loaded only when reading docx
    from docx import Document  # pyright: ignore[reportMissingModuleSource]  # noqa: PLC0415

    # test-seam: re-imported here so monkeypatch on
    # ``mdd.convert.images.write_image`` reaches this call site.
    from mdd.convert.images import write_image  # noqa: PLC0415

    doc: Any = Document(str(src))  # pyright: ignore[reportAny]
    rels = doc.part.rels  # pyright: ignore[reportAny]
    cache: dict[str, Path] = {}
    links: list[str] = []
    for rid in _ordered_image_rids(doc):
        rel: Any = rels.get(rid)  # pyright: ignore[reportAny]
        if rel is None or rel.reltype != _IMAGE_REL_TYPE:  # pyright: ignore[reportAny]
            links.append("")
            continue
        part: Any = rel.target_part  # pyright: ignore[reportAny]
        content_type: str = part.content_type  # pyright: ignore[reportAny]
        fmt = _CONTENT_TYPE_TO_FORMAT.get(content_type, content_type.split("/")[-1])
        blob: bytes = part.blob  # pyright: ignore[reportAny]
        result = write_image(
            attachments_dir, blob, fmt, cache=cache, on_drop=lambda r: dropped.update({r: 1})
        )
        if result is None:
            links.append("")
            continue
        rel_path = f"{attachments_dir.name}/{result.rel_path.as_posix()}"
        links.append(f"![]({rel_path})")
    return links


def _inject_image_links(md: str, links: list[str]) -> str:
    """Replace ``<!-- image -->`` placeholders left by Docling with markdown links.

    Replacements happen in document order. If we collected more links than
    Docling left placeholders, the extras are appended at the end so the
    image isn't silently dropped from the markdown — the plan's
    acceptance criterion is "image present, link somewhere in the same
    block" rather than perfect positioning.
    """
    placeholder = "<!-- image -->"
    remaining = md
    consumed = 0
    out_parts: list[str] = []
    for link in links:
        idx = remaining.find(placeholder)
        if idx == -1:
            break
        replacement = link or "<!-- image dropped -->"
        out_parts.append(remaining[:idx])
        out_parts.append(replacement)
        remaining = remaining[idx + len(placeholder) :]
        consumed += 1
    out_parts.append(remaining)
    out = "".join(out_parts)

    leftover_links = [link for link in links[consumed:] if link]
    if leftover_links:
        out = out.rstrip() + "\n\n" + "\n\n".join(leftover_links) + "\n"
    return out


def _format_metadata(meta: dict[str, str]) -> str:
    lines: list[str] = []
    for key, value in meta.items():
        if key in _MULTI_VALUE_KEYS:
            names = [n.strip() for n in value.split(",") if n.strip()]
            lines.append(f"- **{key}:**")
            lines.extend(f"  - {name}" for name in names)
        elif key in _DOC_META_KEYS:
            lines.append(f"- {key}: {value}")
        else:
            lines.append(f"- **{key}:** {value}")
    return "\n".join(lines)


def _write_docx(src: Path, dst: Path) -> Path | None:
    """Convert a .docx file to markdown, writing to dst.

    Returns the attachments dir if any embedded images were extracted,
    or ``None`` if the doc had no images (so callers can leave
    ``ConvertResult.attachments_dir`` as ``None``).

    Raises :class:`mdd.convert.CorruptSourceError` when *src* is empty
    or is not a valid Office Open XML package. The dispatcher routes
    this to a soft-skip counter (issue #129).
    """
    # lazy: python-docx pulls lxml; loaded only when reading docx
    from docx.opc.exceptions import (  # pyright: ignore[reportMissingModuleSource]  # noqa: PLC0415
        PackageNotFoundError,
    )

    # Cheap pre-check: an empty file cannot be a ZIP container.
    if src.stat().st_size == 0:
        raise CorruptSourceError(f"{src}: empty file (0 bytes)")

    try:
        title = extract_title(src)
        meta = extract_metadata(src)
        body = convert_body(src)
    except PackageNotFoundError as exc:
        raise CorruptSourceError(f"{src}: not a valid docx package ({exc})") from exc

    attachments_dir = dst.parent / (dst.stem + "-attachments")
    dropped: Counter[str] = Counter()
    links = _extract_docx_images(src, attachments_dir, dropped)
    body = _inject_image_links(body, links)
    if dropped:
        breakdown = ", ".join(
            f"{fmt}×{count}"
            for fmt, count in sorted(dropped.items(), key=lambda kv: (-kv[1], kv[0]))
        )
        log.warning(
            "%s: dropped %d embedded image(s) that could not be loaded (unsupported format: %s).",
            src.name,
            sum(dropped.values()),
            breakdown,
        )

    if meta:
        body = strip_leading_table(body)

    parts: list[str] = []
    if title:
        parts.append(f"# {title}")
        parts.append("")
    if meta:
        parts.append(_format_metadata(meta))
        parts.append("")
    body_stripped = body.lstrip("\n")
    parts.append(body_stripped)
    content = "\n".join(parts)

    dst.parent.mkdir(parents=True, exist_ok=True)
    tmp = dst.with_suffix(".md.tmp")
    tmp.write_text(content, encoding="utf-8")
    tmp.rename(dst)

    if attachments_dir.is_dir() and any(attachments_dir.iterdir()):
        return attachments_dir
    return None


class DocxConverter:
    """Convert .docx files to .docx.md (Docling + python-docx)."""

    extensions: tuple[str, ...] = (".docx",)
    output_suffix: str = ".md"

    def convert(self, src: Path, *, dest: Path | None = None) -> ConvertResult:
        if dest is None:
            dest = src.parent / (src.name + self.output_suffix)
        warnings: list[str] = []
        try:
            attachments_dir = _write_docx(src, dest)
        except CorruptSourceError:
            # Soft-skip signal; dispatcher will log it as [SKIP] not [ERROR].
            raise
        except Exception:
            log.exception("error converting %s", src)
            raise
        meta_dict: dict[str, object] = {}
        title = extract_title(src)
        if title:
            meta_dict["title"] = title
        raw_meta = extract_metadata(src)
        meta_dict.update(raw_meta)
        return ConvertResult(
            output_path=dest,
            attachments_dir=attachments_dir,
            metadata=meta_dict,
            warnings=warnings,
        )
