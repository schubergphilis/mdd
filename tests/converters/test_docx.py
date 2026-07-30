"""Tests for mdd.converters.docx (DocxConverter and helpers)."""

from typing import TYPE_CHECKING, Any

from mdd.converters.docx import (
    DocxConverter,
    extract_metadata,
    extract_title,
    strip_boilerplate_lines,
    strip_boilerplate_sections,
    strip_empty_headings,
    strip_leading_table,
    strip_toc,
)

if TYPE_CHECKING:
    from pathlib import Path

    import pytest


def _docx_with_paragraph(tmp_path: Path, name: str, text: str) -> Path:
    from docx import Document  # pyright: ignore[reportMissingModuleSource]

    doc: Any = Document()  # pyright: ignore[reportAny]
    doc.add_paragraph(text)
    p = tmp_path / name
    doc.save(str(p))
    return p


def _docx_with_table(
    tmp_path: Path, name: str, rows: int, cols: int, cells: dict[tuple[int, int], str]
) -> Path:
    from docx import Document  # pyright: ignore[reportMissingModuleSource]

    doc: Any = Document()  # pyright: ignore[reportAny]
    table: Any = doc.add_table(rows=rows, cols=cols)  # pyright: ignore[reportAny]
    for (r, c), val in cells.items():
        table.cell(r, c).text = val
    p = tmp_path / name
    doc.save(str(p))
    return p


class TestStripEmptyHeadings:
    def test_removes_empty_headings(self) -> None:
        lines = ["## ", "### ", "# ", "## Real heading"]
        assert strip_empty_headings(lines) == ["## Real heading"]

    def test_keeps_nonempty_headings(self) -> None:
        lines = ["# Title", "## Section", "content"]
        assert strip_empty_headings(lines) == lines

    def test_strips_artifact_headings(self) -> None:
        lines = ["########### 1 Agenda", "## Real heading"]
        result = strip_empty_headings(lines)
        assert "########### 1 Agenda" not in result
        assert "## Real heading" in result


class TestStripToc:
    def test_removes_word_toc(self) -> None:
        lines = [
            "## Table of Contents",
            "* Introduction 1",
            "* Background 2",
            "## Introduction",
            "Real content here.",
        ]
        result = strip_toc(lines)
        assert "## Table of Contents" not in result
        assert "## Introduction" in result

    def test_passthrough_when_no_toc(self) -> None:
        lines = ["# Hello", "Some text.", "## Section"]
        assert strip_toc(lines) == lines


class TestStripLeadingTable:
    def test_removes_leading_table(self) -> None:
        md = "| A | B |\n|---|---|\n| x | y |\n\n## Content\nHello."
        assert strip_leading_table(md) == "## Content\nHello."

    def test_no_table_passthrough(self) -> None:
        md = "## Content\nHello."
        assert strip_leading_table(md) == md


class TestStripBoilerplateSections:
    def test_strips_meeting_logistics(self) -> None:
        lines = [
            "## 1 Key points",
            "- Content",
            "## 2 Meeting Logistics",
            "- Boilerplate",
            "## 3 Actions",
        ]
        result = strip_boilerplate_sections(lines)
        assert "## 2 Meeting Logistics" not in result
        assert "## 3 Actions" in result


class TestStripBoilerplateLines:
    def test_strips_jira_bullet(self) -> None:
        lines = [
            "## Actions",
            "    - Actions are on the [TEC JIRA board](https://jira.example.com) .",
            "- Real item",
        ]
        result = strip_boilerplate_lines(lines)
        assert not any("TEC JIRA board" in ln for ln in result)
        assert "- Real item" in result


class TestExtractTitle:
    def test_returns_empty_for_plain_docx(self, tmp_path: Path) -> None:
        p = _docx_with_paragraph(tmp_path, "plain.docx", "Hello")
        assert extract_title(p) == ""

    def test_returns_core_properties_title(self, tmp_path: Path) -> None:
        from docx import Document  # pyright: ignore[reportMissingModuleSource]

        doc: Any = Document()  # pyright: ignore[reportAny]
        doc.core_properties.title = "My Report"
        p = tmp_path / "titled.docx"
        doc.save(str(p))
        assert extract_title(p) == "My Report"


class TestExtractMetadata:
    def test_returns_empty_when_no_table(self, tmp_path: Path) -> None:
        p = _docx_with_paragraph(tmp_path, "test.docx", "Hello world")
        assert extract_metadata(p) == {}

    def test_extracts_from_two_column_table(self, tmp_path: Path) -> None:
        p = _docx_with_table(
            tmp_path,
            "meta.docx",
            rows=3,
            cols=2,
            cells={
                (0, 0): "Date",
                (0, 1): "2025-10-13",
                (1, 0): "Location",
                (1, 1): "Hybrid",
                (2, 0): "Participants",
                (2, 1): "Alice, Bob",
            },
        )
        result = extract_metadata(p)
        assert result["Date"] == "2025-10-13"
        assert result["Location"] == "Hybrid"

    def test_skips_tables_over_20_rows(self, tmp_path: Path) -> None:
        p = _docx_with_table(
            tmp_path, "big.docx", rows=21, cols=2, cells={(0, 0): "Date", (0, 1): "2025-01-01"}
        )
        assert extract_metadata(p) == {}


class TestDocxConverter:
    def test_convert_returns_result(self, tmp_path: Path) -> None:
        from unittest.mock import patch

        p = _docx_with_paragraph(tmp_path, "hello.docx", "Hello World")
        with patch("mdd.converters.docx.convert_body", return_value="# Hello"):
            result = DocxConverter().convert(p)
        assert result.output_path == p.parent / "hello.docx.md"
        assert result.output_path.exists()

    def test_convert_respects_dest(self, tmp_path: Path) -> None:
        from unittest.mock import patch

        p = _docx_with_paragraph(tmp_path, "hello.docx", "Hello World")
        dest = tmp_path / "out" / "hello.docx.md"
        with patch("mdd.converters.docx.convert_body", return_value="# Hello"):
            result = DocxConverter().convert(p, dest=dest)
        assert result.output_path == dest
        assert dest.exists()

    def test_convert_attachments_dir_none_when_absent(self, tmp_path: Path) -> None:
        from unittest.mock import patch

        p = _docx_with_paragraph(tmp_path, "hello.docx", "Hello")
        with patch("mdd.converters.docx.convert_body", return_value="# Hello"):
            result = DocxConverter().convert(p)
        # No attachments dir is created for plain docx
        assert result.attachments_dir is None


class TestDroppedImageSuppression:
    """Regression: docling's per-image Pillow warnings are swallowed and summarized."""

    def test_dropped_image_warning_suppressed_and_summarized(
        self, tmp_path: Path, capsys: Any, caplog: pytest.LogCaptureFixture
    ) -> None:
        import logging

        from mdd.converters.docx import (
            _SuppressDroppedImageWarnings,  # pyright: ignore[reportPrivateUsage]
        )

        src = tmp_path / "Templated.docx"
        src.write_bytes(b"placeholder")  # not actually opened — context manager only

        docling_log = logging.getLogger("docling.backend.msword_backend")
        mdd_log = logging.getLogger("mdd.converters.docx")
        mdd_log.addHandler(caplog.handler)
        try:
            with (
                caplog.at_level(logging.WARNING, logger="mdd.converters.docx"),
                _SuppressDroppedImageWarnings(src),
            ):
                # Simulate docling emitting two image-load warnings.
                docling_log.warning(
                    "Warning: image cannot be loaded by Pillow: cannot identify image file <bytes>"
                )
                docling_log.warning("Warning: DrawingML image cannot be loaded by Pillow")
                # An unrelated warning must still propagate.
                docling_log.warning("Some other docling issue")
        finally:
            mdd_log.removeHandler(caplog.handler)

        captured = capsys.readouterr()
        # Per-image noise from docling is gone from stderr…
        assert "cannot be loaded by Pillow" not in captured.err
        # …replaced with one summary mentioning the file and count, emitted via logging.
        msgs = " ".join(r.getMessage() for r in caplog.records)
        assert "Templated.docx" in msgs
        assert "dropped 2 embedded image(s)" in msgs

    def test_no_summary_when_no_images_dropped(self, tmp_path: Path, capsys: Any) -> None:
        from mdd.converters.docx import (
            _SuppressDroppedImageWarnings,  # pyright: ignore[reportPrivateUsage]
        )

        src = tmp_path / "Clean.docx"
        src.write_bytes(b"placeholder")
        with _SuppressDroppedImageWarnings(src):
            pass

        captured = capsys.readouterr()
        assert "dropped" not in captured.err


class TestDocxConverterEndToEnd:
    def test_convert_real_docx(self, tmp_path: Path) -> None:
        p = _docx_with_paragraph(tmp_path, "real.docx", "Hello integration")
        result = DocxConverter().convert(p)
        assert result.output_path.exists()
        content = result.output_path.read_text(encoding="utf-8")
        assert content.strip()


def _docx_with_images(tmp_path: Path, name: str, image_count: int) -> Path:
    """Build a tiny .docx with N distinct embedded PNG images."""
    import io as _io

    from docx import Document  # pyright: ignore[reportMissingModuleSource]
    from PIL import Image

    doc: Any = Document()  # pyright: ignore[reportAny]
    doc.add_heading("Doc with pictures", 1)
    for i in range(image_count):
        doc.add_paragraph(f"Para before image {i + 1}")
        img = Image.new("RGB", (32, 32), color=(i * 30, 100, 200))
        buf = _io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        doc.add_picture(buf)
    p = tmp_path / name
    doc.save(str(p))
    return p


class TestDocxImageExtraction:
    """Embedded docx images are written to an attachments dir and linked."""

    def test_two_images_written_and_linked(self, tmp_path: Path) -> None:
        src = _docx_with_images(tmp_path, "pics.docx", image_count=2)
        result = DocxConverter().convert(src)
        attachments = tmp_path / "pics.docx-attachments"

        assert attachments.is_dir()
        files = list(attachments.iterdir())
        assert len(files) == 2

        md = result.output_path.read_text(encoding="utf-8")
        link_count = md.count("![](pics.docx-attachments/")
        assert link_count == 2

        # ConvertResult points at the attachments dir.
        assert result.attachments_dir == attachments

    def test_no_images_no_attachments_dir(self, tmp_path: Path) -> None:
        p = _docx_with_paragraph(tmp_path, "plain.docx", "Just words")
        result = DocxConverter().convert(p)
        assert result.attachments_dir is None
        assert not (tmp_path / "plain.docx-attachments").exists()

    def test_duplicate_images_dedup_within_doc(self, tmp_path: Path) -> None:
        """Two references to the same blob within one doc share one file."""
        import io as _io

        from docx import Document  # pyright: ignore[reportMissingModuleSource]
        from PIL import Image

        doc: Any = Document()  # pyright: ignore[reportAny]
        img = Image.new("RGB", (32, 32), color=(50, 200, 100))
        buf = _io.BytesIO()
        img.save(buf, format="PNG")
        # python-docx writes the blob into a fresh image part each call;
        # the underlying packaging layer may dedup, but to make the test
        # robust regardless we use the SAME BytesIO buffer twice. Even if
        # python-docx writes both parts, the shared writer dedups by hash.
        buf.seek(0)
        doc.add_picture(buf)
        buf.seek(0)
        doc.add_picture(buf)
        src = tmp_path / "dup.docx"
        doc.save(str(src))

        result = DocxConverter().convert(src)
        attachments = tmp_path / "dup.docx-attachments"
        assert attachments.is_dir()
        # One unique blob -> one file on disk.
        assert len(list(attachments.iterdir())) == 1
        # Two image references in the markdown.
        md = result.output_path.read_text(encoding="utf-8")
        assert md.count("![](dup.docx-attachments/") == 2


class TestDocxImageInjection:
    """Unit tests for the placeholder → link injection helper."""

    def test_replaces_placeholders_in_order(self) -> None:
        from mdd.converters.docx import (
            _inject_image_links,  # pyright: ignore[reportPrivateUsage]
        )

        md = "Before\n<!-- image -->\nMiddle\n<!-- image -->\nAfter"
        out = _inject_image_links(md, ["![](a.png)", "![](b.png)"])
        assert "![](a.png)" in out
        assert "![](b.png)" in out
        # First placeholder gets first link, second gets second.
        assert out.index("![](a.png)") < out.index("![](b.png)")
        # Placeholders are gone.
        assert "<!-- image -->" not in out

    def test_more_links_than_placeholders_appends(self) -> None:
        from mdd.converters.docx import (
            _inject_image_links,  # pyright: ignore[reportPrivateUsage]
        )

        md = "Body without image placeholders"
        out = _inject_image_links(md, ["![](a.png)"])
        assert "![](a.png)" in out

    def test_empty_link_emits_dropped_marker(self) -> None:
        from mdd.converters.docx import (
            _inject_image_links,  # pyright: ignore[reportPrivateUsage]
        )

        md = "x\n<!-- image -->\ny"
        out = _inject_image_links(md, [""])
        assert "<!-- image dropped -->" in out


class TestWriteDocxCorruptSource:
    """_write_docx raises CorruptSourceError for empty/garbage input."""

    def test_zero_byte_docx_raises_corrupt(self, tmp_path: Path) -> None:
        import pytest

        from mdd.convert import CorruptSourceError
        from mdd.converters.docx import _write_docx  # pyright: ignore[reportPrivateUsage]

        src = tmp_path / "empty.docx"
        src.write_bytes(b"")
        dst = tmp_path / "empty.docx.md"
        with pytest.raises(CorruptSourceError, match="empty file"):
            _write_docx(src, dst)
        assert not dst.exists()

    def test_not_a_zip_docx_raises_corrupt(self, tmp_path: Path) -> None:
        import pytest

        from mdd.convert import CorruptSourceError
        from mdd.converters.docx import _write_docx  # pyright: ignore[reportPrivateUsage]

        src = tmp_path / "tiny.docx"
        src.write_bytes(b"not a docx\n")
        dst = tmp_path / "tiny.docx.md"
        with pytest.raises(CorruptSourceError, match="not a valid docx package"):
            _write_docx(src, dst)
        assert not dst.exists()

    def test_corrupt_docx_via_converter_class_propagates(self, tmp_path: Path) -> None:
        """DocxConverter.convert re-raises CorruptSourceError without printing [ERROR]."""
        import pytest

        from mdd.convert import CorruptSourceError

        src = tmp_path / "broken.docx"
        src.write_bytes(b"")
        dst = tmp_path / "broken.docx.md"
        with pytest.raises(CorruptSourceError):
            DocxConverter().convert(src, dest=dst)


class TestSbpDocxHeuristicsGate:
    """The site-specific metadata + boilerplate heuristics can be turned off."""

    def _reset(self) -> None:
        from mdd.converters.docx import set_sbp_docx_heuristics

        set_sbp_docx_heuristics(enabled=True)

    def test_extract_metadata_returns_empty_when_disabled(self, tmp_path: Path) -> None:
        from mdd.converters.docx import set_sbp_docx_heuristics

        doc = _docx_with_paragraph(tmp_path, "note.docx", "hello")
        try:
            set_sbp_docx_heuristics(enabled=False)
            assert extract_metadata(doc) == {}
        finally:
            self._reset()

    def test_boilerplate_kept_when_disabled(self) -> None:
        from mdd.converters import docx as docx_mod

        postprocess = docx_mod._postprocess  # pyright: ignore[reportPrivateUsage]
        set_sbp_docx_heuristics = docx_mod.set_sbp_docx_heuristics

        md = (
            "# Notes\n\n## Meeting Logistics\n\n"
            "- Actions are on the [board](x)\n\n## Real\n\nbody\n"
        )
        try:
            # Enabled (default): the Meeting Logistics section is stripped.
            assert "Meeting Logistics" not in postprocess(md)
            # Disabled: the section survives untouched.
            set_sbp_docx_heuristics(enabled=False)
            out = postprocess(md)
            assert "Meeting Logistics" in out
            assert "Actions are on the" in out
        finally:
            self._reset()
